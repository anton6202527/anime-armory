#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
export.py — 章节 md 合并 → txt / docx / 大纲 md / n2d-script 目录结构。

用法:
    python3 export.py <作品根> [--formats txt,docx,outline,n2d] [--title <书名>]

缺省 --formats 取 _meta.json 里的 outputs。
缺省 --title = "<原作名>-<配角名>外传"。

依赖: python-docx（仅 --formats 含 docx 或 n2d 时）
"""
import argparse
import json
import os
import re
import shutil
import sys
from datetime import date

CHAPTER_FILE_RE = re.compile(r"^第(\d+)章\.md$")
META_LINE_RE = re.compile(r"^<!--\s*meta:.*-->\s*$")
H1_RE = re.compile(r"^#\s+第\s*\d+\s*章\s*[《<]?([^》>]*)[》>]?\s*$")


def load_meta(project_root):
    with open(os.path.join(project_root, "_meta.json"), "r", encoding="utf-8") as f:
        return json.load(f)


def collect_chapters(project_root):
    """读章节/第NN章.md，返回 [(idx, title, body_lines), ...] 按 idx 升序。"""
    chap_dir = os.path.join(project_root, "章节")
    if not os.path.isdir(chap_dir):
        return []
    items = []
    for fname in os.listdir(chap_dir):
        m = CHAPTER_FILE_RE.match(fname)
        if not m:
            continue
        idx = int(m.group(1))
        with open(os.path.join(chap_dir, fname), "r", encoding="utf-8") as f:
            lines = f.read().splitlines()
        # 解析 H1 提取标题
        title = ""
        body_start = 0
        for i, ln in enumerate(lines):
            mh = H1_RE.match(ln.strip())
            if mh:
                title = mh.group(1).strip()
                body_start = i + 1
                break
        body_lines = []
        for ln in lines[body_start:]:
            if META_LINE_RE.match(ln):
                continue
            body_lines.append(ln)
        # 去开头空行
        while body_lines and not body_lines[0].strip():
            body_lines.pop(0)
        items.append((idx, title, body_lines))
    items.sort(key=lambda x: x[0])
    return items


def total_chars(chapters):
    total = 0
    for _, title, body in chapters:
        total += len(title)
        for ln in body:
            total += len(ln.strip())
    return total


def write_txt(out_path, meta, chapters, title):
    total = total_chars(chapters)
    provenance = [
        f"# spinoff_of: {meta['source_title']}",
        f"# spinoff_character: {meta['spinoff_character']}",
        f"# mode: {meta['mode']}",
        f"# chapters: {len(chapters)}",
        f"# chars: {total}",
        f"# rights_status: {meta['rights_status']}",
        f"# generated: {date.today().isoformat()}",
        f"# tool: novel-spinoff",
    ]
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(provenance) + "\n\n")
        for idx, title_c, body in chapters:
            head = f"第{idx}章 {title_c}".rstrip()
            f.write(head + "\n\n")
            for ln in body:
                f.write(ln + "\n")
            f.write("\n")


def write_docx(out_path, meta, chapters, title):
    try:
        from docx import Document
    except ImportError:
        print("[err] 缺依赖：pip install python-docx", file=sys.stderr)
        sys.exit(2)
    doc = Document()
    total = total_chars(chapters)
    # provenance 块（普通段落）
    provenance_lines = [
        f"原作：{meta['source_title']}",
        f"视角：{meta['spinoff_character']}",
        f"模式：{meta['mode']}    规模：{meta['scale']}    章数：{len(chapters)}    字数：{total}",
        f"版权状态：{meta['rights_status']}    生成日期：{date.today().isoformat()}",
        f"工具：novel-spinoff",
    ]
    for ln in provenance_lines:
        doc.add_paragraph(ln)
    doc.add_paragraph("")
    for idx, title_c, body in chapters:
        doc.add_heading(f"第{idx}章 {title_c}".rstrip(), level=1)
        # 把连续非空行视为一段
        para_buf = []
        for ln in body:
            if ln.strip():
                para_buf.append(ln.strip())
            else:
                if para_buf:
                    doc.add_paragraph("".join(para_buf))
                    para_buf = []
        if para_buf:
            doc.add_paragraph("".join(para_buf))
    doc.save(out_path)


def write_outline(out_path, project_root, meta, chapters):
    """读 设定/章纲.md，剥内部注释，加现章节统计。"""
    outline_src = os.path.join(project_root, "设定", "章纲.md")
    if os.path.exists(outline_src):
        with open(outline_src, "r", encoding="utf-8") as f:
            content = f.read()
        # 剥引用块注释行（以 `> ` 开头的元说明）
        kept = []
        for ln in content.splitlines():
            if ln.lstrip().startswith("> "):
                continue
            kept.append(ln)
        cleaned = "\n".join(kept).strip()
    else:
        cleaned = f"# 章纲 — {meta['spinoff_character']}外传\n\n（章纲未填）"

    summary = (
        f"\n\n---\n\n"
        f"_共 {len(chapters)} 章，{total_chars(chapters)} 字。"
        f"原作《{meta['source_title']}》，模式 {meta['mode']}。_\n"
    )
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(cleaned + summary)


def write_n2d(n2d_root, docx_path, title):
    """铺 n2d-script 友好的目录结构。"""
    novel_dir = os.path.join(n2d_root, "小说")
    os.makedirs(novel_dir, exist_ok=True)
    shutil.copy(docx_path, os.path.join(novel_dir, f"{title}.docx"))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("project_root", help="作品根（init_project.py 建的那个）")
    ap.add_argument("--formats", default=None,
                    help="逗号分隔，可含 txt,docx,outline,n2d；缺省 = _meta.json.outputs")
    ap.add_argument("--title", default=None, help="缺省 = <原作名>-<配角名>外传")
    args = ap.parse_args()

    project_root = os.path.abspath(args.project_root)
    if not os.path.isdir(project_root):
        print(f"[err] 找不到作品根：{project_root}", file=sys.stderr)
        sys.exit(2)

    meta = load_meta(project_root)
    formats = (args.formats.split(",") if args.formats else meta.get("outputs", []))
    formats = [f.strip() for f in formats if f.strip()]
    title = (
        args.title
        or meta.get("title")
        or f"{meta['source_title']}-{meta['spinoff_character']}外传"
    )

    chapters = collect_chapters(project_root)
    if not chapters:
        print("[err] 章节/ 下没有 第NN章.md，先写章节再导出", file=sys.stderr)
        sys.exit(2)

    out_dir = os.path.join(project_root, "导出")
    os.makedirs(out_dir, exist_ok=True)

    paths = {}
    docx_path = None

    if "txt" in formats:
        p = os.path.join(out_dir, f"{title}.txt")
        write_txt(p, meta, chapters, title)
        paths["txt"] = p
    if "docx" in formats or "n2d" in formats:
        p = os.path.join(out_dir, f"{title}.docx")
        write_docx(p, meta, chapters, title)
        paths["docx"] = p
        docx_path = p
    if "outline" in formats:
        p = os.path.join(out_dir, "大纲.md")
        write_outline(p, project_root, meta, chapters)
        paths["outline"] = p
    if "n2d" in formats:
        n2d_root = os.path.join(out_dir, "n2d-script")
        write_n2d(n2d_root, docx_path, title)
        paths["n2d"] = n2d_root

    print(f"[ok] 导出完成：{len(chapters)} 章, {total_chars(chapters)} 字")
    for k, v in paths.items():
        print(f"     {k:<8} → {v}")
    if "n2d" in paths:
        print(f"[next] 进 n2d-script：python3 skills/n2d-script/scripts/split_novel.py "
              f"\"{os.path.join(paths['n2d'], '小说', title + '.docx')}\"")


if __name__ == "__main__":
    main()
