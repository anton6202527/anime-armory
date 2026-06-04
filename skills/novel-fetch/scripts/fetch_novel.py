#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fetch_novel.py — 给定章节目录页 URL，联网抓取公版小说全文，输出 .txt + .docx。
（合法性：默认只抓公版/开放授权来源；付费墙站直接拒抓。详见 ../SKILL.md 与 ../references/sources.md。）

用法:
    python3 fetch_novel.py <目录页URL> --name "<书名>" [--out <作品根>]

选项:
    --name 书名      输出文件名与标题（必填）
    --out  作品根    输出到 <作品根>/小说/；缺省 = 写小说/<书名>/
    --source auto|gutenberg|wikisource|generic   抓取引擎（默认 auto 探测）
    --i-have-rights  对非公版/通用兜底 URL 声明你有权使用（跳过合法性确认）

依赖: requests beautifulsoup4 trafilatura python-docx
"""
import argparse
import os
import re
import sys
from urllib.parse import quote, unquote, urljoin

# 依赖：import 名 -> pip 安装名
_DEP_INSTALL_NAME = {
    "requests": "requests",
    "bs4": "beautifulsoup4",
    "trafilatura": "trafilatura",
    "docx": "python-docx",
}


def _detect_have():
    have = set()
    for mod in _DEP_INSTALL_NAME:
        try:
            __import__(mod)
            have.add(mod)
        except ImportError:
            pass
    return have


def missing_deps(have=None):
    """返回缺失依赖的 pip 安装名列表（按 _DEP_INSTALL_NAME 顺序）。"""
    if have is None:
        have = _detect_have()
    return [install for mod, install in _DEP_INSTALL_NAME.items() if mod not in have]


# 已知付费墙/反爬站：命中直接拒抓（不替用户规避）
PAYWALL_DOMAINS = (
    "qidian.com", "fanqienovel.com", "jjwxc.net", "jjwxc.com",
    "zongheng.com", "17k.com", "hongxiu.com", "yuewen.com",
    "ciweimao.com", "faloo.com", "readnovel.com",
)


def _host(url):
    m = re.match(r"^[a-z]+://([^/]+)", url.strip(), re.I)
    return (m.group(1) if m else url).lower()


def is_paywalled(url):
    host = _host(url)
    return any(host == d or host.endswith("." + d) for d in PAYWALL_DOMAINS)


def detect_source(url):
    host = _host(url)
    if "gutenberg.org" in host or "gutendex.com" in host:
        return "gutenberg"
    if "wikisource.org" in host:
        return "wikisource"
    return "generic"


def assemble_text(chapters):
    """把 [{title, body}] 合并成纯文本：每章一行 `第N章 标题`，空行，正文。
    章节标题格式通用、利于后续按章拆分。"""
    blocks = []
    for i, ch in enumerate(chapters, 1):
        title = (ch.get("title") or "").strip()
        body = (ch.get("body") or "").strip()
        blocks.append(f"第{i}章 {title}\n\n{body}")
    return "\n\n".join(blocks) + "\n"


_CHAPTER_RE = re.compile(r"^第\s*[0-9零一二三四五六七八九十百千两]+\s*[章回节卷]")


def _provenance_lines(prov):
    return [
        "# === 抓取来源信息 (provenance) — 注释块，按章拆分时会自动跳过 ===",
        f"# source_url: {prov.get('source_url', '')}",
        f"# fetched: {prov.get('fetched', '')}",
        f"# chapters: {prov.get('chapters', '')}",
        f"# chars: {prov.get('chars', '')}",
        f"# copyright: {prov.get('copyright', '')}",
        "# ================================================================",
    ]


def write_txt(path, text, prov):
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    header = "\n".join(_provenance_lines(prov))
    with open(path, "w", encoding="utf-8") as f:
        f.write(header + "\n\n" + text)


def write_docx(path, text, prov):
    import docx
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc = docx.Document()
    for line in _provenance_lines(prov):
        doc.add_paragraph(line.lstrip("# ").rstrip())
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        first, _, rest = block.partition("\n")
        if _CHAPTER_RE.match(first.strip()):
            doc.add_heading(first.strip(), level=1)
            if rest.strip():
                for para in rest.split("\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
        else:
            for para in block.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
    doc.save(path)


def extract_body(html, url=None):
    """从单页 HTML 提取正文：优先 trafilatura，失败回退 readability，再回退 bs4 取最长 <div>。"""
    try:
        import trafilatura
        extracted = trafilatura.extract(html, url=url, favor_recall=True,
                                         include_comments=False, include_tables=False)
        if extracted and extracted.strip():
            return extracted.strip()
    except ImportError:
        pass
    try:
        from readability import Document as _RDoc
        from bs4 import BeautifulSoup
        summary_html = _RDoc(html).summary()
        return BeautifulSoup(summary_html, "html.parser").get_text("\n").strip()
    except ImportError:
        pass
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    blocks = soup.find_all(["div", "article"])
    best = max(blocks, key=lambda b: len(b.get_text()), default=soup.body or soup)
    return best.get_text("\n").strip()


# 章节链接锚文本特征：含「第…章/回/节/卷」或「序/楔子/尾声/番外」
_CH_LINK_RE = re.compile(
    r"第\s*[0-9零一二三四五六七八九十百千两]+\s*[章回节卷]|楔子|序章|序言|尾声|番外")


def extract_chapter_links(html, base_url):
    """从目录页提取章节链接 [(绝对URL, 标题)]，按出现顺序，去重保序。"""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    out, seen = [], set()
    for a in soup.find_all("a", href=True):
        text = a.get_text(strip=True)
        if not text or not _CH_LINK_RE.search(text):
            continue
        url = urljoin(base_url, a["href"])
        if url in seen:
            continue
        seen.add(url)
        out.append((url, text))
    return out


def http_get(url):
    """单页 GET，返回解码后的 HTML 文本。仅此一处真实联网。"""
    import requests
    resp = requests.get(url, headers={"User-Agent": "novel-fetch/1.0"}, timeout=30)
    resp.raise_for_status()
    if not resp.encoding or resp.encoding.lower() == "iso-8859-1":
        resp.encoding = resp.apparent_encoding
    return resp.text


def fetch_generic(index_url, get=http_get):
    """通用兜底：抓目录页 → 遍历章节页 → 逐章提正文。返回 [{title, body}]。
    逐章打印抓取状态到 stderr。"""
    index_html = get(index_url)
    links = extract_chapter_links(index_html, base_url=index_url)
    if not links:
        raise SystemExit("目录页未发现章节链接；请确认这是章节目录页 URL。")
    chapters = []
    for i, (url, title) in enumerate(links, 1):
        try:
            body = extract_body(get(url), url=url)
            status = "ok" if body else "empty"
        except Exception as e:  # noqa: BLE001 — 单章失败不应中断全书
            body, status = "", f"fail({type(e).__name__})"
        chapters.append({"title": title, "body": body})
        print(f"  [{i}/{len(links)}] {status}: {title}", file=sys.stderr)
    return chapters


_PLAINTEXT_CH_RE = re.compile(
    r"^\s*(Chapter\s+[IVXLCDM0-9]+|CHAPTER\s+[IVXLCDM0-9]+|"
    r"第\s*[0-9零一二三四五六七八九十百千两]+\s*[章回节卷])\b.*$", re.M)


def split_plaintext_chapters(raw):
    """把纯文本（Gutenberg）按 Chapter/第N章 标记切章；无标记则整本一章。"""
    marks = list(_PLAINTEXT_CH_RE.finditer(raw))
    if not marks:
        return [{"title": "正文", "body": raw.strip()}]
    # 从第一个章节标记处开始迭代，intentionally 丢弃前置非章节内容（版权声明等前言）
    chapters = []
    for i, m in enumerate(marks):
        start = m.end()
        end = marks[i + 1].start() if i + 1 < len(marks) else len(raw)
        chapters.append({"title": m.group(0).strip(), "body": raw[start:end].strip()})
    return chapters


def _gutenberg_book_id(url):
    m = re.search(r"/(?:ebooks|books)/(\d+)", url)
    if not m:
        raise SystemExit(f"无法从 URL 解析 Gutenberg 书号: {url}")
    return m.group(1)


def fetch_gutenberg(url, get=http_get, get_json=None):
    """Project Gutenberg：经 gutendex 拿 plain-text 链接 → 下载 → 切章。"""
    import json as _json
    if get_json is None:
        def get_json(u):
            return _json.loads(get(u))
    book_id = _gutenberg_book_id(url)
    meta = get_json(f"https://gutendex.com/books/{book_id}")
    fmts = meta.get("formats", {})
    txt_url = next((v for k, v in fmts.items()
                    if k.startswith("text/plain") and not v.endswith(".zip")), None)
    if not txt_url:
        raise SystemExit("gutendex 未提供纯文本格式；换一本或用 --source generic。")
    return split_plaintext_chapters(get(txt_url))


def fetch_wikisource(page_url, get=http_get):
    """中文维基文库：MediaWiki action=parse 取渲染 HTML → extract_body。
    单页作品 → 一章；多卷请对每卷分页 URL 各跑一次（v1 先支持单页）。"""
    m = re.match(r"^(https?://[^/]+)/wiki/(.+)$", page_url)
    if not m:
        raise SystemExit(f"不是有效的 Wikisource 页面 URL: {page_url}")
    api = m.group(1) + "/w/api.php"
    title = unquote(m.group(2))
    url = (f"{api}?action=parse&prop=text&format=json"
           f"&formatversion=2&page={quote(title)}")
    import json as _json
    data = _json.loads(get(url))
    if "error" in data:
        raise SystemExit(f"Wikisource API 错误: {data['error'].get('info', data['error'])}")
    html = data.get("parse", {}).get("text", "")
    if isinstance(html, dict):  # formatversion<2 兼容
        html = html.get("*", "")
    body = extract_body(html, url=page_url)
    return [{"title": title.split("/")[-1], "body": body}]


def _today():
    """抓取日期（脚本运行时刻）。"""
    import datetime
    return datetime.date.today().isoformat()


def resolve_out_dir(out, name):
    """输出目录 = <作品根>/小说/；缺省作品根 = 写小说/<书名>/。
    若 --out 本身已指向 小说/ 目录，直接用它，避免 小说/小说 双层嵌套。"""
    if out and os.path.basename(out.rstrip("/\\")) == "小说":
        return out
    root = out if out else os.path.join("写小说", name)
    return os.path.join(root, "小说")


def dispatch(url, source="auto", engines=None):
    """按 source 路由到对应抓取引擎，返回 [{title, body}]。"""
    if engines is None:
        engines = {"gutenberg": fetch_gutenberg,
                   "wikisource": fetch_wikisource,
                   "generic": fetch_generic}
    chosen = source if source != "auto" else detect_source(url)
    if chosen not in engines:
        raise SystemExit(f"未知 source: {chosen}")
    return engines[chosen](url)


def main():
    ap = argparse.ArgumentParser(description="联网抓取公版小说全文 → txt + docx")
    ap.add_argument("url", help="章节目录页 / 作品页 URL")
    ap.add_argument("--name", required=True, help="书名（输出文件名与标题）")
    ap.add_argument("--out", default=None, help="作品根；缺省 = 写小说/<书名>/")
    ap.add_argument("--source", default="auto",
                    choices=["auto", "gutenberg", "wikisource", "generic"])
    ap.add_argument("--i-have-rights", action="store_true",
                    help="对非公版/通用兜底 URL 声明有权使用")
    args = ap.parse_args()

    missing = missing_deps()
    if missing:
        sys.exit("缺少依赖，请先安装：pip install " + " ".join(missing))

    if is_paywalled(args.url):
        sys.exit("拒抓：该站为已知付费墙/反爬来源，本工具不替你规避。请改用公版来源。")

    src = args.source if args.source != "auto" else detect_source(args.url)
    if src == "generic" and not args.i_have_rights:
        sys.exit("通用兜底抓取非公版来源需声明授权：确认你有权使用后加 --i-have-rights 重跑。")

    chapters = dispatch(args.url, source=src)
    chapters = [c for c in chapters if c.get("body")]
    if not chapters:
        sys.exit("未抓到任何正文。请检查 URL 是否为章节目录页。")

    text = assemble_text(chapters)
    chars = len(text.replace("\n", ""))
    copyright_note = {"gutenberg": "公版（Project Gutenberg）",
                      "wikisource": "公版（中文维基文库）"}.get(src, "用户声明有权使用")
    prov = {"source_url": args.url, "fetched": _today(), "chapters": len(chapters),
            "chars": chars, "copyright": copyright_note}

    out_dir = resolve_out_dir(args.out, args.name)
    txt_path = os.path.join(out_dir, args.name + ".txt")
    docx_path = os.path.join(out_dir, args.name + ".docx")
    write_txt(txt_path, text, prov)
    write_docx(docx_path, text, prov)

    print(f"完成：{len(chapters)} 章，{chars} 字")
    print(f"  txt : {txt_path}")
    print(f"  docx: {docx_path}")


if __name__ == "__main__":
    main()
