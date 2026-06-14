#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
export.py — novel-* 家族通用导出器（单一真值源）。

章节/第NN章.md 合并 → txt / docx / 大纲 md / n2d-script 目录结构。
被 novel-create / spinoff / rewrite / expand / condense / continue 共用，
各 skill 不再各写一份 export 脚本（旧的 expand.py/condense.py/continue.py 已删除）。

用法:
    python3 export.py <作品根> [--formats txt,docx,outline,n2d] [--title <书名>] [--combine]

缺省 --formats 取 _meta.json 里的 outputs。
缺省 --title 按 kind 推导（见 derive_title）。
--combine 仅用于 novel-continue：原作 + 新章节合一输出（章号续编），输出单个合本 txt。

依赖: python-docx（仅 --formats 含 docx 或 n2d 时）
"""
import argparse
import json
import os
import re
import shutil
import sys
from datetime import date

from contract import ALLOWED_OUTPUT_FORMATS, derive_title
from qa_gate import collect_gate_status, format_gate_status
from report_snapshot import sha256_file, snapshot_chapters
from waivers import append_waiver, make_waiver

CHAPTER_FILE_RE = re.compile(r"^第0*(\d+)章(?:[_ ].*)?\.md$")  # 第N章.md 或 第N章_标题.md
META_LINE_RE = re.compile(r"^<!--\s*meta:.*-->\s*$")
# 章号接受阿拉伯/全角/中文数字；标题可带《》或裸标题
H1_RE = re.compile(r"^#\s+第\s*[0-9０-９一二三四五六七八九十百千零〇两]+\s*章\s*[《<]?([^》>]*)[》>]?\s*$")
ORIG_CHAP_RE = re.compile(r"^\s*第\s*[0-9零一二三四五六七八九十百千两]+\s*[章回节卷]")

def load_meta(project_root):
    with open(os.path.join(project_root, "_meta.json"), "r", encoding="utf-8") as f:
        return json.load(f)


def collect_chapters(project_root):
    """读章节/第NN章.md，返回 [(idx, title, body_lines), ...] 按 idx 升序。"""
    chap_dir = os.path.join(project_root, "章节")
    if not os.path.isdir(chap_dir):
        return []
    items = []
    for fname in os.listdir(chap_dir):
        m = CHAPTER_FILE_RE.match(fname)
        if not m:
            continue
        idx = int(m.group(1))
        with open(os.path.join(chap_dir, fname), "r", encoding="utf-8") as f:
            lines = f.read().splitlines()
        title = ""
        body_start = 0
        for i, ln in enumerate(lines):
            mh = H1_RE.match(ln.strip())
            if mh:
                title = mh.group(1).strip()
                body_start = i + 1
                break
        body_lines = []
        for ln in lines[body_start:]:
            if META_LINE_RE.match(ln):
                continue
            body_lines.append(ln)
        while body_lines and not body_lines[0].strip():
            body_lines.pop(0)
        items.append((idx, title, body_lines))
    items.sort(key=lambda x: x[0])
    return items


def total_chars(chapters):
    total = 0
    for _, title, body in chapters:
        total += len(title)
        for ln in body:
            total += len(ln.strip())
    return total


def count_orig_chapters(orig_txt_path):
    if not os.path.exists(orig_txt_path):
        return 0
    return sum(1 for ln in open(orig_txt_path, encoding="utf-8") if ORIG_CHAP_RE.match(ln))


def _provenance_lines(meta, chapters):
    """txt 头部 provenance —— 各 kind 共用，按存在的字段补充。"""
    kind = meta.get("kind", "spinoff")
    lines = [
        f"# source: {meta.get('source_title') or meta.get('source') or '—'}",
        f"# kind: {kind}",
        f"# chapters: {len(chapters)}",
        f"# chars: {total_chars(chapters)}",
        f"# rights_status: {meta.get('rights_status', '—')}",
        f"# rights_jurisdiction: {meta.get('rights_jurisdiction', '—')}",
        f"# distribution_regions: {','.join(meta.get('distribution_regions') or []) or '—'}",
        f"# generated: {date.today().isoformat()}",
        f"# tool: novel-{kind}",
    ]
    for field, label in (("spinoff_character", "spinoff_character"),
                         ("rewrite_type", "rewrite_type"),
                         ("ratio", "ratio"),
                         ("mode", "mode"),
                         ("direction_chosen", "direction")):
        if meta.get(field) not in (None, ""):
            lines.insert(2, f"# {label}: {meta[field]}")
    return lines


def write_txt(out_path, meta, chapters, title):
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(_provenance_lines(meta, chapters)) + "\n\n")
        for idx, title_c, body in chapters:
            f.write(f"第{idx}章 {title_c}".rstrip() + "\n\n")
            for ln in body:
                f.write(ln + "\n")
            f.write("\n")


def write_combined_txt(out_path, project_root, meta, new_chapters):
    """novel-continue --combine：原作正文 + 新章节（章号续编）合一。"""
    orig_path = os.path.join(project_root, "原作.txt")
    orig_count = count_orig_chapters(orig_path)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(_provenance_lines(meta, new_chapters)) + "\n")
        f.write(f"# combined: True (原作 {orig_count} 章 + 新 {len(new_chapters)} 章)\n\n")
        # 原作正文（跳过头部 # 注释/空行）
        if os.path.exists(orig_path):
            kept, skip = [], True
            for ln in open(orig_path, encoding="utf-8").read().splitlines():
                if skip and (ln.startswith("#") or ln.strip() == ""):
                    continue
                skip = False
                kept.append(ln)
            f.write("\n".join(kept))
            if kept and kept[-1].strip():
                f.write("\n")
        f.write("\n--- 续写新章节 ---\n\n")
        for idx, t, body in new_chapters:
            f.write(f"第{orig_count + idx}章 {t}".rstrip() + "\n\n")
            for ln in body:
                f.write(ln + "\n")
            f.write("\n")
    return orig_count


def write_docx(out_path, meta, chapters, title):
    try:
        from docx import Document
    except ImportError:
        print("[err] 缺依赖：pip install python-docx", file=sys.stderr)
        sys.exit(2)
    doc = Document()
    total = total_chars(chapters)
    kind = meta.get("kind", "spinoff")
    src = meta.get("source_title") or meta.get("source") or "—"
    mode = meta.get("mode") or meta.get("rewrite_type") or kind or "—"
    provenance_lines = [f"原作：{src}"]
    if meta.get("spinoff_character"):
        provenance_lines.append(f"视角：{meta['spinoff_character']}")
    provenance_lines += [
        f"模式：{mode}    规模：{meta.get('scale', '—')}    章数：{len(chapters)}    字数：{total}",
        f"版权状态：{meta.get('rights_status', '—')}    权利辖区：{meta.get('rights_jurisdiction', '—')}    生成日期：{date.today().isoformat()}",
        f"工具：novel-{kind}",
    ]
    for ln in provenance_lines:
        doc.add_paragraph(ln)
    doc.add_paragraph("")
    for idx, title_c, body in chapters:
        doc.add_heading(f"第{idx}章 {title_c}".rstrip(), level=1)
        para_buf = []
        for ln in body:
            if ln.strip():
                para_buf.append(ln.strip())
            else:
                if para_buf:
                    doc.add_paragraph("".join(para_buf))
                    para_buf = []
        if para_buf:
            doc.add_paragraph("".join(para_buf))
    doc.save(out_path)


def write_outline(out_path, project_root, meta, chapters):
    """读 设定/章纲.md，剥内部注释，加现章节统计。"""
    outline_src = os.path.join(project_root, "设定", "章纲.md")
    if os.path.exists(outline_src):
        with open(outline_src, "r", encoding="utf-8") as f:
            content = f.read()
        kept = [ln for ln in content.splitlines() if not ln.lstrip().startswith("> ")]
        cleaned = "\n".join(kept).strip()
    else:
        who = meta.get("spinoff_character") or meta.get("title") or "本作"
        cleaned = f"# 章纲 — {who}\n\n（章纲未填）"
    src = meta.get("source_title") or meta.get("source") or "—"
    mode = meta.get("mode") or meta.get("kind") or meta.get("rewrite_type") or "—"
    summary = (
        f"\n\n---\n\n"
        f"_共 {len(chapters)} 章，{total_chars(chapters)} 字。"
        f"原作《{src}》，模式 {mode}。_\n"
    )
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(cleaned + summary)


def _find_drama_repo_root(start):
    """向上找含『制漫剧/』的仓库根；找不到返回 None。与 split_novel 的作品根定位同源。"""
    d = os.path.abspath(start)
    while True:
        if os.path.isdir(os.path.join(d, "制漫剧")):
            return d
        parent = os.path.dirname(d)
        if parent == d:
            return None
        d = parent


def resolve_n2d_dest(project_root, title, explicit):
    """决定 n2d 交接落点（= 作品根，其下铺 小说/）。返回 (dest, mode)。

    - explicit（--n2d-dest）优先。
    - 否则自动定位 <repo>/制漫剧/<title>/：这样 split_novel 直接吃 小说/<title>.docx
      就会把 n2d 生产树建在正确位置（split 取 小说/ 的父级为作品根），无需 --out。
    - 都不行才回退项目内 导出/n2d-script/（此时 split 必须带 --out，否则会建错位置）。
    """
    if explicit:
        return os.path.abspath(explicit), "explicit"
    repo = _find_drama_repo_root(project_root)
    if repo:
        return os.path.join(repo, "制漫剧", title), "canonical"
    return os.path.join(project_root, "导出", "n2d-script"), "legacy"


def write_n2d(n2d_root, docx_path, title, meta, project_root):
    """铺 n2d-script 友好的目录结构 + 交接清单（_n2d_handoff.json，留痕来源/版权/hash）。返回落地 docx 路径。"""
    novel_dir = os.path.join(n2d_root, "小说")
    os.makedirs(novel_dir, exist_ok=True)
    dest_docx = os.path.join(novel_dir, f"{title}.docx")
    shutil.copy(docx_path, dest_docx)
    handoff = {
        "schema_version": 1,
        "source_novel_project": os.path.basename(project_root.rstrip("/\\")),
        "source_novel_path": project_root,
        "source_title": meta.get("source_title") or meta.get("source") or "",
        "title": title,
        "kind": meta.get("kind", ""),
        "rights_status": meta.get("rights_status", ""),
        "rights_jurisdiction": meta.get("rights_jurisdiction", ""),
        "rights_basis": meta.get("rights_basis", ""),
        "source_license_url": meta.get("source_license_url", ""),
        "rights_covered_regions": meta.get("rights_covered_regions", []),
        "distribution_regions": meta.get("distribution_regions", []),
        "requires_region_rights_review": meta.get("requires_region_rights_review", False),
        "docx": f"{title}.docx",
        "docx_sha256": sha256_file(dest_docx),
        "exported": date.today().isoformat(),
    }
    with open(os.path.join(novel_dir, "_n2d_handoff.json"), "w", encoding="utf-8") as f:
        json.dump(handoff, f, ensure_ascii=False, indent=2)

    # Asset-Aware Extraction
    asset_registry = {"characters": {}, "props": {}, "vfx": {}, "locations": {}, "outfits": {}}
    tag_pattern = re.compile(r"\[(CHAR|PROP|VFX|LOC|OUTFIT)_([^\]]+)\]")
    ch_dir = os.path.join(project_root, "章节")
    if os.path.exists(ch_dir):
        for fname in sorted(os.listdir(ch_dir)):
            if fname.endswith(".md"):
                with open(os.path.join(ch_dir, fname), "r", encoding="utf-8") as f:
                    for match in tag_pattern.finditer(f.read()):
                        kind, name = match.group(1), match.group(2)
                        key = f"{kind}_{name}"
                        if kind == "CHAR" and key not in asset_registry["characters"]:
                            asset_registry["characters"][key] = {"id": key, "name": name, "source_chapter": fname}
                        elif kind == "PROP" and key not in asset_registry["props"]:
                            asset_registry["props"][key] = {"id": key, "name": name, "source_chapter": fname}
                        elif kind == "VFX" and key not in asset_registry["vfx"]:
                            asset_registry["vfx"][key] = {"id": key, "name": name, "source_chapter": fname}
                        elif kind == "LOC" and key not in asset_registry["locations"]:
                            asset_registry["locations"][key] = {"id": key, "name": name, "source_chapter": fname}
                        elif kind == "OUTFIT" and key not in asset_registry["outfits"]:
                            asset_registry["outfits"][key] = {"id": key, "name": name, "source_chapter": fname}
    
    if any(asset_registry.values()):
        with open(os.path.join(novel_dir, "asset_registry_preflight.json"), "w", encoding="utf-8") as f:
            json.dump(asset_registry, f, ensure_ascii=False, indent=2)

    return dest_docx


def _rel(project_root, path):
    return os.path.relpath(os.path.abspath(path), project_root).replace(os.sep, "/")


def export_waiver_scope(project_root, gate_status, formats, *, combine=False):
    chapter_snapshot = snapshot_chapters(project_root, mode="export")
    reports = []
    for report in gate_status.get("reports") or []:
        path = report.get("path")
        if path and os.path.exists(path):
            reports.append({
                "kind": report.get("kind"),
                "path": _rel(project_root, path),
                "sha256": sha256_file(path),
            })
    return {
        "output_mode": "combine" if combine else "export",
        "formats": ["combine"] if combine else sorted(formats or []),
        "chapter_count": len(chapter_snapshot.get("files") or []),
        "source_aggregate_hash": chapter_snapshot.get("aggregate_hash") or "",
        "blocker_ids": sorted(str(b.get("id") or "") for b in gate_status.get("blockers") or []),
        "reports": reports,
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("project_root", help="作品根（init_project.py 建的那个）")
    ap.add_argument("--formats", default=None,
                    help="逗号分隔，可含 txt,docx,outline,n2d；缺省 = _meta.json.outputs")
    ap.add_argument("--title", default=None, help="缺省按 kind 推导")
    ap.add_argument("--combine", action="store_true",
                    help="novel-continue 合本：原作 + 新章节合一（章号续编）")
    ap.add_argument("--ignore-qa-gate", action="store_true",
                    help="强制导出：忽略 review/score 阻断报告（只用于人工明确要求）")
    ap.add_argument("--n2d-dest", default=None,
                    help="n2d 交接落点（制漫剧作品根，其下铺 小说/）。缺省自动取 <repo>/制漫剧/<书名>/，"
                         "让 split_novel 无需 --out 即建在正确位置；仅 n2d 格式生效")
    args = ap.parse_args()

    project_root = os.path.abspath(args.project_root)
    if not os.path.isdir(project_root):
        print(f"[err] 找不到作品根：{project_root}", file=sys.stderr)
        sys.exit(2)

    meta = load_meta(project_root)
    title = args.title or derive_title(meta)
    chapters = collect_chapters(project_root)
    if not chapters:
        print("[err] 章节/ 下没有 第NN章.md / 第NN章_标题.md，先写章节再导出", file=sys.stderr)
        sys.exit(2)

    if args.combine:
        formats = ["combine"]
    else:
        formats = (args.formats.split(",") if args.formats else meta.get("outputs", []))
        formats = [f.strip() for f in formats if f.strip()]
        if not formats:
            print("[err] 未指定导出格式：请传 --formats txt,docx,outline,n2d，"
                  "或在 _meta.json 写 outputs。", file=sys.stderr)
            sys.exit(2)
        unknown = sorted(set(formats) - set(ALLOWED_OUTPUT_FORMATS))
        if unknown:
            print(f"[err] 未知导出格式：{','.join(unknown)}；可用："
                  f"{','.join(ALLOWED_OUTPUT_FORMATS)}", file=sys.stderr)
            sys.exit(2)

    gate_status = collect_gate_status(project_root, require_review_report=True, export_formats=formats)
    if gate_status["blocking"] and not args.ignore_qa_gate:
        print(format_gate_status(gate_status), file=sys.stderr)
        print("[err] QA gate 阻断导出；按报告回流修改，或人工确认后加 --ignore-qa-gate。", file=sys.stderr)
        sys.exit(1)
    if gate_status["blocking"] and args.ignore_qa_gate:
        waiver = make_waiver(
            "ignore_qa_gate",
            reason="explicit --ignore-qa-gate during export",
            affected_gate="export_qa_gate",
            source="novel-craft/scripts/export.py",
            details={"blockers": gate_status["blockers"]},
            scope=export_waiver_scope(project_root, gate_status, formats, combine=args.combine),
        )
        waiver["risk"] = "本次导出绕过了 review/score QA gate；仅对当前章节 hash、阻断项和导出格式有效。"
        log_path = append_waiver(project_root, waiver)
        print(format_gate_status(gate_status), file=sys.stderr)
        print(f"[warn] 已强制导出并记录 QA gate 豁免：{log_path}", file=sys.stderr)
    elif gate_status.get("warnings"):
        print(format_gate_status(gate_status), file=sys.stderr)

    out_dir = os.path.join(project_root, "导出")
    os.makedirs(out_dir, exist_ok=True)

    # --combine：续写合本，单 txt 输出（其它格式忽略）
    if args.combine:
        src = meta.get("source_title") or title
        p = os.path.join(out_dir, f"{src}-合本.txt")
        orig_count = write_combined_txt(p, project_root, meta, chapters)
        print(f"[ok] 合本：原作 {orig_count} 章 + 新 {len(chapters)} 章 / {total_chars(chapters)} 新字 → {p}")
        print(f"     新章节起始编号 = 第 {orig_count + 1} 章")
        return

    paths = {}
    docx_path = None
    if "txt" in formats:
        p = os.path.join(out_dir, f"{title}.txt")
        write_txt(p, meta, chapters, title)
        paths["txt"] = p
    if "docx" in formats or "n2d" in formats:
        p = os.path.join(out_dir, f"{title}.docx")
        write_docx(p, meta, chapters, title)
        paths["docx"] = p
        docx_path = p
    if "outline" in formats:
        p = os.path.join(out_dir, "大纲.md")
        write_outline(p, project_root, meta, chapters)
        paths["outline"] = p
    n2d_mode = n2d_docx = None
    if "n2d" in formats:
        dest, n2d_mode = resolve_n2d_dest(project_root, title, args.n2d_dest)
        n2d_docx = write_n2d(dest, docx_path, title, meta, project_root)
        paths["n2d"] = dest

    print(f"[ok] 导出完成：{len(chapters)} 章, {total_chars(chapters)} 字")
    for k, v in paths.items():
        print(f"     {k:<8} → {v}")
    if "n2d" in paths:
        if n2d_mode == "legacy":
            print(f"[warn] 未找到含『制漫剧/』的仓库根，n2d 交接回退项目内：{paths['n2d']}")
            print(f"[next] 进 n2d-script（须带 --out 指定剧名，否则会建错位置）：\n"
                  f"       python3 skills/n2d-script/scripts/split_novel.py \"{n2d_docx}\" --out 制漫剧/{title}")
        else:
            print(f"[ok] n2d 交接已落入作品根：{paths['n2d']}（小说/ + _n2d_handoff.json）")
            print(f"[next] 进 n2d-script（小说在 制漫剧/<剧名>/小说/ 下，split 自动取父级为作品根，无需 --out）：\n"
                  f"       python3 skills/n2d-script/scripts/split_novel.py \"{n2d_docx}\"")


if __name__ == "__main__":
    main()
