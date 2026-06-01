#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fetch_novel.py — 给定章节目录页 URL，联网抓取公版小说全文，输出 .txt + .docx。
（合法性：默认只抓公版/开放授权来源；付费墙站直接拒抓。详见 ../SKILL.md 与 ../references/sources.md。）

用法:
    python3 fetch_novel.py <目录页URL> --name "<书名>" [--out <作品根>]

选项:
    --name 书名      输出文件名与标题（必填）
    --out  作品根    输出到 <作品根>/小说/；缺省 = artifacts/<书名>/
    --source auto|gutenberg|wikisource|generic   抓取引擎（默认 auto 探测）
    --i-have-rights  对非公版/通用兜底 URL 声明你有权使用（跳过合法性确认）

依赖: requests beautifulsoup4 trafilatura python-docx
"""
import argparse
import os
import re
import sys

# 依赖：import 名 -> pip 安装名
_DEP_INSTALL_NAME = {
    "requests": "requests",
    "bs4": "beautifulsoup4",
    "trafilatura": "trafilatura",
    "docx": "python-docx",
}


def _detect_have():
    have = set()
    for mod in _DEP_INSTALL_NAME:
        try:
            __import__(mod)
            have.add(mod)
        except ImportError:
            pass
    return have


def missing_deps(have=None):
    """返回缺失依赖的 pip 安装名列表（按 _DEP_INSTALL_NAME 顺序）。"""
    if have is None:
        have = _detect_have()
    return [install for mod, install in _DEP_INSTALL_NAME.items() if mod not in have]


# 已知付费墙/反爬站：命中直接拒抓（不替用户规避）
PAYWALL_DOMAINS = (
    "qidian.com", "fanqienovel.com", "jjwxc.net", "jjwxc.com",
    "zongheng.com", "17k.com", "hongxiu.com", "yuewen.com",
    "ciweimao.com", "faloo.com", "readnovel.com",
)


def _host(url):
    m = re.match(r"^[a-z]+://([^/]+)", url.strip(), re.I)
    return (m.group(1) if m else url).lower()


def is_paywalled(url):
    host = _host(url)
    return any(host == d or host.endswith("." + d) for d in PAYWALL_DOMAINS)


def detect_source(url):
    host = _host(url)
    if "gutenberg.org" in host or "gutendex.com" in host:
        return "gutenberg"
    if "wikisource.org" in host:
        return "wikisource"
    return "generic"


def assemble_text(chapters):
    """把 [{title, body}] 合并成符合 split_novel.py 输入约定的纯文本：
    每章一行 `第N章 标题`，空行，正文。"""
    blocks = []
    for i, ch in enumerate(chapters, 1):
        title = (ch.get("title") or "").strip()
        body = (ch.get("body") or "").strip()
        blocks.append(f"第{i}章 {title}\n\n{body}")
    return "\n\n".join(blocks) + "\n"


_CHAPTER_RE = re.compile(r"^第\s*[0-9零一二三四五六七八九十百千两]+\s*[章回节卷]")


def _provenance_lines(prov):
    return [
        "# === 抓取来源信息 (provenance) — split_novel.py 会自动跳过本块 ===",
        f"# source_url: {prov.get('source_url', '')}",
        f"# fetched: {prov.get('fetched', '')}",
        f"# chapters: {prov.get('chapters', '')}",
        f"# chars: {prov.get('chars', '')}",
        f"# copyright: {prov.get('copyright', '')}",
        "# ================================================================",
    ]


def write_txt(path, text, prov):
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    header = "\n".join(_provenance_lines(prov))
    with open(path, "w", encoding="utf-8") as f:
        f.write(header + "\n\n" + text)


def write_docx(path, text, prov):
    import docx
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc = docx.Document()
    for line in _provenance_lines(prov):
        doc.add_paragraph(line.lstrip("# ").rstrip())
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        first, _, rest = block.partition("\n")
        if _CHAPTER_RE.match(first.strip()):
            doc.add_heading(first.strip(), level=1)
            if rest.strip():
                for para in rest.split("\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
        else:
            for para in block.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
    doc.save(path)
