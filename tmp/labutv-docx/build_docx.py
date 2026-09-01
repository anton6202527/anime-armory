#!/usr/bin/env python3
"""Build the polished LabuTV Desktop guide DOCX from the repository Markdown."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "LabuTV-Electron-App.md"
OUTPUT = ROOT / "tmp" / "labutv-docx" / "LabuTV-Electron-App-产品与使用指南.docx"
IMAGES = SOURCE.parent / "app-screenshots"
DOC_TOOLS = Path(
    "/Users/wesley/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.826.12353/skills/documents/scripts"
)
sys.path.insert(0, str(DOC_TOOLS))

from table_geometry import (  # noqa: E402
    DEFAULT_CELL_MARGINS_DXA,
    apply_table_geometry,
    column_widths_from_weights,
    section_content_width_dxa,
)


BODY_FONT = "Arial"
CJK_FONT = "Hiragino Sans GB"
CODE_FONT = "Courier New"
NAVY = "183B56"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "2E9E91"
GOLD = "C4962E"
INK = "202A33"
MID = "5E6974"
LIGHT = "F4F7FA"
TABLE_HEADER = "E8EEF5"
TABLE_ALT = "F8FAFC"
BORDER = "CBD6E2"
WHITE = "FFFFFF"


def set_run_font(run, *, size=None, bold=None, italic=None, color=None, name=None):
    """Set Latin and CJK fonts explicitly so Chinese remains portable."""

    font_name = name or BODY_FONT
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_fonts.set(qn("w:cs"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, name=BODY_FONT, cjk=CJK_FONT, size=None, color=None, bold=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), cjk)
    style._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_border(paragraph, *, side="bottom", color=BLUE, size="12", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    borders.append(edge)


def set_paragraph_shading(paragraph, fill=LIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_keep(paragraph, *, next_=False, lines=True):
    paragraph.paragraph_format.keep_with_next = next_
    paragraph.paragraph_format.keep_together = lines
    paragraph.paragraph_format.widow_control = True


def add_hyperlink(paragraph, text, url, *, bold=False):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), BODY_FONT)
    r_fonts.set(qn("w:hAnsi"), BODY_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.append(r_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def public_url(target: str) -> str:
    if target.startswith(("http://", "https://", "mailto:")):
        return target
    resolved = (SOURCE.parent / target).resolve()
    try:
        rel = resolved.relative_to(ROOT).as_posix()
    except ValueError:
        return target
    return f"https://github.com/anton6202527/anime-armory/blob/main/{rel}"


INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|``.+?``|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def add_inline(paragraph, text: str, *, default_bold=False, size=None, color=None):
    """Render the small Markdown subset used by the source guide."""

    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, bold=default_bold, color=color)
        token = match.group(0)
        if token.startswith("**"):
            inner = token[2:-2]
            run = paragraph.add_run(inner)
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("``"):
            inner = token[2:-2]
            run = paragraph.add_run(inner)
            set_run_font(run, size=(size or 10.5) - 0.5, bold=default_bold, color=DARK_BLUE, name=CODE_FONT)
            set_run_shading(run, "EEF3F8")
        elif token.startswith("`"):
            inner = token[1:-1]
            run = paragraph.add_run(inner)
            set_run_font(run, size=(size or 10.5) - 0.5, bold=default_bold, color=DARK_BLUE, name=CODE_FONT)
            set_run_shading(run, "EEF3F8")
        else:
            label, target = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, public_url(target), bold=default_bold)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=default_bold, color=color)


def set_run_shading(run, fill):
    r_pr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    r_pr.append(shd)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, placeholder, end])
    set_run_font(run, size=8.5, color=MID)


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    set_style_font(caption, size=9, color=MID)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_together = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    set_style_font(code, name=CODE_FONT, cjk=CJK_FONT, size=9, color="28323C")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.05
    code.paragraph_format.keep_together = True

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    set_style_font(callout, size=10, color=DARK_BLUE)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.08)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.18

    if "TOC Line" not in styles:
        toc = styles.add_style("TOC Line", WD_STYLE_TYPE.PARAGRAPH)
    else:
        toc = styles["TOC Line"]
    set_style_font(toc, size=10.5, color=INK)
    toc.paragraph_format.space_after = Pt(4)
    toc.paragraph_format.left_indent = Inches(0.12)
    toc.paragraph_format.keep_together = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def configure_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("LabuTV Desktop  ·  Electron 产品与使用指南")
    set_run_font(run, size=8.5, color=MID)
    set_paragraph_border(p, side="bottom", color=BORDER, size="4", space="3")

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.space_before = Pt(3)
    run = p.add_run("LabuTV  ·  2026-08-28")
    set_run_font(run, size=8.5, color=MID)
    p.add_run("\t")
    add_page_field(p)


def append_num_pr(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


class Numbering:
    def __init__(self, document):
        self.root = document.part.numbering_part.element
        abstract_ids = [
            int(el.get(qn("w:abstractNumId")))
            for el in self.root.findall(qn("w:abstractNum"))
        ]
        num_ids = [int(el.get(qn("w:numId"))) for el in self.root.findall(qn("w:num"))]
        self.next_abstract = max(abstract_ids, default=0) + 1
        self.next_num = max(num_ids, default=0) + 1
        self.abstract = {
            "bullet": self._add_abstract("bullet", "•"),
            "decimal": self._add_abstract("decimal", "%1."),
        }

    def _add_abstract(self, fmt, text):
        abstract_id = self.next_abstract
        self.next_abstract += 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        abstract.append(lvl)
        self.root.append(abstract)
        return abstract_id

    def new(self, kind):
        num_id = self.next_num
        self.next_num += 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(self.abstract[kind]))
        num.append(abstract_num_id)
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
        self.root.append(num)
        return num_id


def add_cover(document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("PRODUCT & USER GUIDE")
    set_run_font(run, size=10, bold=True, color=GOLD)
    set_paragraph_border(p, side="bottom", color=GOLD, size="16", space="8")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("LabuTV Desktop")
    set_run_font(run, size=34, bold=True, color=NAVY)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Electron 产品与使用指南")
    set_run_font(run, size=22, bold=True, color=BLUE)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("本地 AI 内容生产工作台 · 从自然语言需求到可检查、可返修、可验收的最终产物")
    set_run_font(run, size=11, color=MID)

    picture = IMAGES / "labutv-01-home.png"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(picture), width=Inches(6.25))
    shape._inline.docPr.set("descr", "LabuTV Desktop 首页，展示自然语言输入、技能选择、工作区和创作线入口")
    shape._inline.docPr.set("name", "LabuTV Desktop 首页")
    p.paragraph_format.space_after = Pt(12)
    set_keep(p, lines=True)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    add_inline(p, "适用对象  希望使用 LabuTV 管理本地 AI 内容生产的创作者、工作室与维护者", size=9.5, color=MID)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_inline(p, "文档版本  2026-08-28", size=9.5, color=MID)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, "文档定位  产品介绍、安装说明、界面导览、n2d 制作流程与安全边界", size=9.5, color=MID)

    document.add_page_break()


def add_contents(document, headings):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("阅读导航")
    set_run_font(run, size=10, bold=True, color=GOLD)

    h = document.add_paragraph(style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    add_inline(h, "目录")

    p = document.add_paragraph(style="Callout")
    set_paragraph_shading(p, "EEF6F5")
    set_paragraph_border(p, side="left", color=TEAL, size="18", space="8")
    add_inline(
        p,
        "快速路径：先看第 2 节下载安装；第 4–11 节熟悉界面；第 12–13 节理解 n2d 的一键制作、状态、哈希与完成定义。",
        size=10,
    )

    for heading in headings:
        p = document.add_paragraph(style="TOC Line")
        run = p.add_run(heading)
        set_run_font(run, size=10.5, color=INK)
        set_paragraph_border(p, side="bottom", color="E6EBF0", size="2", space="2")

    p = document.add_paragraph(style="Callout")
    set_paragraph_shading(p, "FFF8E8")
    set_paragraph_border(p, side="left", color=GOLD, size="18", space="8")
    add_inline(
        p,
        "版本提示：本文截图来自 2026-08-28 当前开发版；公开 Electron 安装包发布于 2026-07-23，品牌名和部分能力可能早于本文截图。请以 GitHub Release 的 tag、说明与资产列表判断实际版本。",
        size=10,
    )
    document.add_page_break()


def add_picture(document, source, alt, figure_no):
    path = (SOURCE.parent / source).resolve()
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as image:
        px_w, px_h = image.size
    max_width = 6.25
    max_height = 7.0
    width = min(max_width, max_height * px_w / px_h)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("name", f"图 {figure_no} · {alt[:40]}")
    set_keep(p, next_=True, lines=True)

    caption = document.add_paragraph(style="Figure Caption")
    add_inline(caption, f"图 {figure_no} · {alt}", size=9, color=MID)
    return figure_no + 1


def split_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def table_weights(headers):
    count = len(headers)
    header_text = " ".join(headers)
    if count == 2:
        if "skill" in header_text.lower():
            return [1.55, 4.95]
        return [1.7, 4.8]
    if count == 3:
        if "创作线" in header_text:
            return [1.0, 0.95, 4.55]
        if "文件" in header_text and "下载" in header_text:
            return [1.9, 2.8, 1.8]
        if "操作" in header_text:
            return [2.45, 1.9, 2.15]
        return [1.5, 2.1, 2.9]
    return [1.0] * count


def add_table(document, rows):
    if not rows:
        return
    count = len(rows[0])
    if any(len(row) != count for row in rows):
        raise ValueError(f"ragged Markdown table: {rows[:2]}")
    table = document.add_table(rows=len(rows), cols=count)
    table.autofit = False
    table.alignment = 0
    table.style = "Table Grid"
    for row_idx, values in enumerate(rows):
        row = table.rows[row_idx]
        set_no_split(row)
        if row_idx == 0:
            set_repeat_table_header(row)
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_shading(cell, TABLE_HEADER if row_idx == 0 else (TABLE_ALT if row_idx % 2 == 0 else WHITE))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.12
            p.paragraph_format.widow_control = True
            add_inline(p, value, default_bold=row_idx == 0, size=9.2, color=NAVY if row_idx == 0 else INK)
    content_width = section_content_width_dxa(document.sections[-1])
    widths = column_widths_from_weights(table_weights(rows[0]), content_width)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=content_width,
        indent_dxa=DEFAULT_CELL_MARGINS_DXA["start"],
        cell_margins_dxa=DEFAULT_CELL_MARGINS_DXA,
    )
    trailing = document.add_paragraph()
    trailing.paragraph_format.space_after = Pt(0)
    trailing.paragraph_format.line_spacing = 0.5


def add_code_block(document, lines):
    for idx, line in enumerate(lines):
        p = document.add_paragraph(style="Code Block")
        p.paragraph_format.keep_with_next = idx < len(lines) - 1
        if idx == 0:
            p.paragraph_format.space_before = Pt(5)
        if idx == len(lines) - 1:
            p.paragraph_format.space_after = Pt(7)
        set_paragraph_shading(p, "F0F3F6")
        if idx == 0:
            set_paragraph_border(p, side="top", color=BORDER, size="4", space="2")
        if idx == len(lines) - 1:
            set_paragraph_border(p, side="bottom", color=BORDER, size="4", space="2")
        run = p.add_run(line or " ")
        set_run_font(run, size=9, color="28323C", name=CODE_FONT)


def add_list(document, items, kind, numbering):
    num_id = numbering.new(kind)
    for item in items:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.widow_control = True
        append_num_pr(p, num_id)
        add_inline(p, item, size=11)


def render_markdown(document, lines):
    numbering = Numbering(document)
    figure_no = 1
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1)) - 1
            text = heading_match.group(2)
            p = document.add_paragraph(style=f"Heading {level}")
            add_inline(p, text)
            if level == 1:
                set_paragraph_border(p, side="bottom", color=BLUE, size="6", space="4")
            i += 1
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if image_match:
            figure_no = add_picture(document, image_match.group(2), image_match.group(1), figure_no)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            raw_rows = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                raw_rows.append(split_table_row(lines[i]))
                i += 1
            rows = [raw_rows[0]] + [row for row in raw_rows[2:]]
            add_table(document, rows)
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            add_code_block(document, code_lines)
            continue

        if re.match(r"^-\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^-\s+", lines[i].strip()):
                items.append(re.sub(r"^-\s+", "", lines[i].strip()))
                i += 1
            add_list(document, items, "bullet", numbering)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            add_list(document, items, "decimal", numbering)
            continue

        if stripped.startswith(">"):
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = document.add_paragraph(style="Callout")
            set_paragraph_shading(p, "FFF8E8")
            set_paragraph_border(p, side="left", color=GOLD, size="18", space="8")
            add_inline(p, " ".join(quote), size=10)
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate:
                break
            if (
                re.match(r"^(#{2,4})\s+", lines[i])
                or candidate.startswith(("![", "|", "```", ">"))
                or re.match(r"^-\s+", candidate)
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            i += 1
        p = document.add_paragraph()
        add_inline(p, " ".join(paragraph_lines), size=11)


def add_document_settings(document):
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)


def build():
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    headings = [
        match.group(1)
        for line in source_lines
        if (match := re.match(r"^##\s+(.+)$", line))
    ]
    first_section_index = next(i for i, line in enumerate(source_lines) if line.startswith("## "))

    document = Document()
    configure_styles(document)
    for section in document.sections:
        configure_section(section)
        configure_header_footer(section)
    add_document_settings(document)

    props = document.core_properties
    props.title = "LabuTV Desktop（Electron）产品与使用指南"
    props.subject = "LabuTV Electron 桌面端产品介绍、使用说明与 n2d 制漫剧流程"
    props.author = "LabuTV"
    props.last_modified_by = "LabuTV"
    props.keywords = "LabuTV, Electron, AI 内容生产, n2d, 制漫剧"
    props.comments = "由仓库内公开产品文档转换；版本 2026-08-28。"

    add_cover(document)
    add_contents(document, headings)
    render_markdown(document, source_lines[first_section_index:])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
